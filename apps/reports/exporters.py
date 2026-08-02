"""
Turns an answer + references into:
  1. A Markdown string (source of truth, shown in the UI)
  2. A PDF file (via markdown -> HTML -> xhtml2pdf)
  3. A DOCX file (via python-docx, with a real References heading)
"""
import io
import markdown as md_lib
from docx import Document
from docx.shared import Pt
from xhtml2pdf import pisa


def build_markdown(question: str, answer: str, references: list[dict]) -> str:
    lines = [f"# Research Report\n", f"**Question:** {question}\n", "---\n", answer, "\n---\n", "## References\n"]
    for ref in references:
        lines.append(f"[{ref['index']}] {ref['title']} — {ref['url']}\n")
    return "\n".join(lines)


def markdown_to_pdf_bytes(markdown_text: str) -> bytes:
    html_body = md_lib.markdown(markdown_text, extensions=["extra", "sane_lists"])
    html_doc = f"""
    <html>
    <head>
    <style>
        body {{ font-family: Helvetica, Arial, sans-serif; font-size: 11px; line-height: 1.5; color: #222; }}
        h1 {{ font-size: 18px; border-bottom: 2px solid #333; padding-bottom: 6px; }}
        h2 {{ font-size: 14px; margin-top: 20px; color: #444; }}
        a {{ color: #1a5fb4; }}
    </style>
    </head>
    <body>{html_body}</body>
    </html>
    """
    buf = io.BytesIO()
    result = pisa.CreatePDF(src=html_doc, dest=buf)
    if result.err:
        raise RuntimeError("PDF generation failed (xhtml2pdf reported an error).")
    return buf.getvalue()


def markdown_to_docx_bytes(question: str, answer: str, references: list[dict]) -> bytes:
    doc = Document()

    title = doc.add_heading("Research Report", level=1)

    doc.add_paragraph().add_run(f"Question: {question}").bold = True
    doc.add_paragraph("")

    # naive markdown-ish rendering: paragraphs + bracketed citations kept inline
    for para in answer.split("\n"):
        if para.strip():
            doc.add_paragraph(para.strip())

    doc.add_heading("References", level=2)
    for ref in references:
        p = doc.add_paragraph()
        p.add_run(f"[{ref['index']}] ").bold = True
        run = p.add_run(f"{ref['title']} — {ref['url']}")
        run.font.size = Pt(10)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
